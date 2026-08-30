"""Parsers that normalize PDF, Word, Excel, Markdown and text into sections."""

from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Callable

import pymupdf
from docx import Document as WordDocument
from openpyxl import load_workbook

from config import Settings, get_settings
from .models import DocumentUnit, ParsedDocument


def _clean_text(text: str) -> str:
    text = text.replace("\x00", " ").replace("\r\n", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _stable_doc_id(path: Path) -> str:
    digest = hashlib.sha256()
    digest.update(path.name.encode("utf-8"))
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()[:24]


class MultiFormatParser:
    def __init__(self, settings: Settings | None = None):
        self.settings = settings or get_settings()
        self._parsers: dict[str, Callable[[Path], list[DocumentUnit]]] = {
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".xlsx": self._parse_xlsx,
            ".md": self._parse_markdown,
            ".txt": self._parse_text,
        }

    def parse(self, file_path: str | Path) -> ParsedDocument:
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix not in self._parsers:
            raise ValueError(f"不支持的文件格式：{suffix or '(无扩展名)'}")
        if not path.is_file():
            raise FileNotFoundError(path)
        if path.stat().st_size > self.settings.max_upload_mb * 1024 * 1024:
            raise ValueError(f"文件超过 {self.settings.max_upload_mb} MB 限制：{path.name}")

        units = [unit for unit in self._parsers[suffix](path) if unit.text.strip()]
        if not units:
            raise ValueError(f"没有从 {path.name} 提取到可索引文本")
        return ParsedDocument(
            doc_id=_stable_doc_id(path),
            source=path.name,
            file_type=suffix,
            units=units,
        )

    @staticmethod
    def _parse_pdf(path: Path) -> list[DocumentUnit]:
        units: list[DocumentUnit] = []
        with pymupdf.open(path) as pdf:
            for index, page in enumerate(pdf, start=1):
                text = _clean_text(page.get_text("text"))
                if text:
                    units.append(
                        DocumentUnit(
                            text=text,
                            heading=f"第 {index} 页",
                            locator=f"page:{index}",
                            metadata={"page": index},
                        )
                    )
        return units

    @staticmethod
    def _parse_docx(path: Path) -> list[DocumentUnit]:
        document = WordDocument(path)
        units: list[DocumentUnit] = []
        heading_stack: list[str] = []
        buffer: list[str] = []

        def flush() -> None:
            text = _clean_text("\n".join(buffer))
            if text:
                units.append(DocumentUnit(text=text, heading=" > ".join(heading_stack)))
            buffer.clear()

        for paragraph in document.paragraphs:
            text = _clean_text(paragraph.text)
            if not text:
                continue
            style_name = (paragraph.style.name or "").lower()
            match = re.search(r"heading\s*(\d+)|标题\s*(\d+)", style_name)
            if match:
                flush()
                level = int(match.group(1) or match.group(2) or 1)
                heading_stack[:] = heading_stack[: level - 1]
                heading_stack.append(text)
            else:
                buffer.append(text)
        flush()

        for table_index, table in enumerate(document.tables, start=1):
            rows = []
            for row in table.rows:
                values = [_clean_text(cell.text) for cell in row.cells]
                if any(values):
                    rows.append(" | ".join(values))
            if rows:
                units.append(
                    DocumentUnit(
                        text="\n".join(rows),
                        heading=f"表格 {table_index}",
                        locator=f"table:{table_index}",
                    )
                )
        return units

    @staticmethod
    def _parse_xlsx(path: Path) -> list[DocumentUnit]:
        workbook = load_workbook(path, read_only=True, data_only=True)
        units: list[DocumentUnit] = []
        try:
            for sheet in workbook.worksheets:
                rows = list(sheet.iter_rows(values_only=True))
                if not rows:
                    continue
                headers = [str(value).strip() if value is not None else f"列{index + 1}" for index, value in enumerate(rows[0])]
                lines: list[str] = []
                start_row = 2
                for row_number, row in enumerate(rows[1:], start=2):
                    pairs = [
                        f"{headers[index]}: {value}"
                        for index, value in enumerate(row)
                        if index < len(headers) and value not in (None, "")
                    ]
                    if pairs:
                        lines.append(f"第{row_number}行 | " + " | ".join(pairs))
                    if len(lines) == 30:
                        units.append(
                            DocumentUnit(
                                text="\n".join(lines),
                                heading=f"工作表：{sheet.title}",
                                locator=f"sheet:{sheet.title};rows:{start_row}-{row_number}",
                            )
                        )
                        lines = []
                        start_row = row_number + 1
                if lines:
                    end_row = start_row + len(lines) - 1
                    units.append(
                        DocumentUnit(
                            text="\n".join(lines),
                            heading=f"工作表：{sheet.title}",
                            locator=f"sheet:{sheet.title};rows:{start_row}-{end_row}",
                        )
                    )
        finally:
            workbook.close()
        return units

    @staticmethod
    def _parse_markdown(path: Path) -> list[DocumentUnit]:
        units: list[DocumentUnit] = []
        headings: list[str] = []
        buffer: list[str] = []

        def flush() -> None:
            text = _clean_text("\n".join(buffer))
            if text:
                units.append(DocumentUnit(text=text, heading=" > ".join(headings)))
            buffer.clear()

        for line in path.read_text(encoding="utf-8").splitlines():
            match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if match:
                flush()
                level = len(match.group(1))
                headings[:] = headings[: level - 1]
                headings.append(match.group(2).strip())
            else:
                buffer.append(line)
        flush()
        return units

    @staticmethod
    def _parse_text(path: Path) -> list[DocumentUnit]:
        text = _clean_text(path.read_text(encoding="utf-8-sig"))
        return [DocumentUnit(text=text, heading=path.stem)] if text else []
