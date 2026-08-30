"""Generate clearly-labelled synthetic enterprise documents for the demo."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.lib import colors


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "sample_data" / "generated"


def generate_policy_docx() -> Path:
    path = OUTPUT / "星云科技员工差旅与报销制度.docx"
    document = Document()
    document.add_heading("星云科技员工差旅与报销制度（模拟）", 0)
    document.add_paragraph("版本：V2.1；生效日期：2026年7月1日。本制度及其中数据均为项目演示所用的虚构内容。")
    sections = {
        "一、出差申请": [
            "员工应在出发前至少2个工作日通过OA提交出差申请。",
            "跨省出差由直属负责人和部门负责人两级审批；海外出差还需总经理审批。",
        ],
        "二、住宿标准": [
            "一线城市住宿上限为每人每晚600元，其他城市为每人每晚450元。",
            "因展会或重大活动导致超标时，须在报销单中说明原因，并由部门负责人追加审批；未经追加审批的超标部分由个人承担。",
        ],
        "三、交通与餐补": [
            "高铁行程原则上购买二等座；飞行时间超过4小时可申请经济舱。",
            "国内出差餐补为每人每天100元，已由接待方提供餐食的对应餐次不重复补贴。",
        ],
        "四、报销时限与凭证": [
            "员工应在出差结束后10个工作日内提交报销。",
            "报销必须附合法电子发票或纸质发票、行程单及审批记录；单张发票金额超过1000元时需附支付凭证。",
        ],
    }
    for heading, paragraphs in sections.items():
        document.add_heading(heading, level=1)
        for value in paragraphs:
            document.add_paragraph(value)
    document.save(path)
    return path


def generate_manual_pdf() -> Path:
    path = OUTPUT / "星云智能会议终端产品手册.pdf"
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    body = ParagraphStyle("ChineseBody", parent=styles["BodyText"], fontName="STSong-Light", fontSize=10.5, leading=17)
    title = ParagraphStyle("ChineseTitle", parent=styles["Title"], fontName="STSong-Light", fontSize=20, alignment=TA_CENTER)
    heading = ParagraphStyle("ChineseHeading", parent=styles["Heading1"], fontName="STSong-Light", fontSize=15, leading=22)
    story = [
        Paragraph("星云智能会议终端产品手册（模拟）", title),
        Spacer(1, 8 * mm),
        Paragraph("本手册中的公司、型号、参数和服务政策均为项目演示所用的虚构内容。", body),
        Paragraph("1. 产品概览", heading),
        Paragraph("NX-MEET-S 面向20人以内会议室，支持1080P视频、双麦克风阵列和有线投屏。NX-MEET-PRO 面向50人以内会议室，支持4K视频、六麦克风阵列、无线投屏和双屏输出。", body),
        Paragraph("2. 接口与网络", heading),
        Paragraph("NX-MEET-PRO 提供2个HDMI输出、1个HDMI输入、2个USB 3.0接口、千兆以太网口，并支持Wi-Fi 6。首次部署建议使用有线网络完成激活。", body),
        PageBreak(),
        Paragraph("3. 安装步骤", heading),
        Paragraph("连接显示器和摄像头后接通电源；设备启动后选择语言和网络；输入企业激活码；完成麦克风与扬声器测试；最后在管理后台绑定会议室。激活失败时先检查系统时间和443端口。", body),
        Paragraph("4. 故障排查", heading),
        Paragraph("画面卡顿时依次检查可用带宽、网络丢包和HDMI线缆。无声音时检查默认音频设备、静音状态和USB连接。状态灯连续红色闪烁表示设备温度过高，应断电并检查通风。", body),
        Paragraph("5. 售后说明", heading),
        Paragraph("硬件质保期以销售合同和产品价格表为准。人为损坏、非授权拆机及不可抗力导致的损坏不属于免费质保范围。", body),
    ]
    SimpleDocTemplate(str(path), pagesize=A4, rightMargin=20 * mm, leftMargin=20 * mm).build(story)
    return path


def generate_price_xlsx() -> Path:
    path = OUTPUT / "星云产品与服务价格表.xlsx"
    workbook = Workbook()
    products = workbook.active
    products.title = "产品价格"
    products.append(["SKU", "产品名称", "含税标价（元）", "标准质保（年）", "适用场景"])
    products.append(["NX-MEET-S", "星云会议终端标准版", 3999, 2, "20人以内会议室"])
    products.append(["NX-MEET-PRO", "星云会议终端专业版", 6999, 3, "50人以内会议室"])
    products.append(["NX-CAM-4K", "星云4K智能摄像头", 1299, 2, "视频采集"])
    services = workbook.create_sheet("服务价格")
    services.append(["服务代码", "服务名称", "年费（元）", "响应时间", "说明"])
    services.append(["SV-BASIC", "基础支持", 0, "2个工作日", "工作日9:00-18:00在线支持"])
    services.append(["SV-PRO", "专业支持", 2999, "4小时", "7×12小时远程支持"])
    services.append(["SV-ONSITE", "现场支持", 6999, "下一个工作日", "每年含2次现场服务"])
    note = workbook.create_sheet("说明")
    note.append(["声明"])
    note.append(["本表中的公司、型号、价格和政策均为项目演示所用的虚构数据。"])
    workbook.save(path)
    return path


def generate_security_md() -> Path:
    path = OUTPUT / "星云科技账号与信息安全制度.md"
    path.write_text(
        """# 星云科技账号与信息安全制度（模拟）

本制度及其中数据均为项目演示所用的虚构内容。版本 V1.3，2026年6月15日生效。

## 密码与多因素认证

公司账号密码长度不得少于12位，并至少包含大写字母、小写字母、数字和特殊字符中的三类。管理员账号和远程访问账号必须启用多因素认证。密码不得通过邮件或即时通讯明文传递。

## 账号生命周期

新员工账号由直属负责人发起申请，信息技术部在入职前一个工作日完成开通。员工离职时，人力资源部应在离职生效前通知信息技术部；信息技术部须在离职生效后2小时内停用账号。

## 数据分级与外发

数据分为公开、内部、机密三级。机密数据外发必须经数据所有者审批，并使用公司批准的加密通道。禁止将内部或机密数据上传到个人网盘。

## 安全事件报告

发现账号异常登录、设备丢失或疑似数据泄露后，应在30分钟内联系信息技术部，并在2小时内提交安全事件单。不得自行删除相关日志。
""",
        encoding="utf-8",
    )
    return path


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    paths = [generate_policy_docx(), generate_manual_pdf(), generate_price_xlsx(), generate_security_md()]
    print("Generated synthetic documents:")
    for path in paths:
        print(f"- {path}")


if __name__ == "__main__":
    main()
