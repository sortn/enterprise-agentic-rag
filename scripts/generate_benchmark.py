"""Generate the frozen benchmark_v1 corpus and evaluation datasets.

The benchmark is fact-first: every rendered document clause originates from a
structured fact with a stable clause code. Ground-truth child/parent IDs are
linked after ingestion by ``link_ground_truth.py``.
"""

from __future__ import annotations

from collections import Counter, defaultdict
from datetime import date, timedelta
import hashlib
import json
from pathlib import Path
import random

from docx import Document
from openpyxl import Workbook
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
BENCHMARK_DIR = ROOT / "evaluation" / "benchmark_v1"
CORPUS_DIR = BENCHMARK_DIR / "corpus"
SEED = 20260827
SCOPES = ("总部", "华东区", "华南区", "西南区", "北方区")
FORMATS = (".md", ".docx", ".pdf", ".xlsx", ".txt")
CATEGORY_ORDER = (
    "exact_term",
    "semantic_paraphrase",
    "numeric_table",
    "parent_context",
    "version_conflict",
    "rewrite_needed",
)


DOMAINS = (
    {
        "code": "HR",
        "name": "人力资源",
        "title": "员工考勤与发展管理办法",
        "subjects": ("年假申请", "调休登记", "培训报名", "转岗申请", "试用期评估"),
        "sections": ("申请与审批", "时限要求", "费用标准", "异常处理", "版本说明"),
        "roles": ("直属负责人", "人力资源经理"),
        "portals": ("HR自助平台", "人才发展门户", "考勤中心"),
        "conditions": ("审批人临时离岗", "系统考勤记录缺失", "培训计划发生冲突"),
        "actions": ("提交书面说明", "补充考勤证据"),
    },
    {
        "code": "FIN",
        "name": "财务",
        "title": "费用报销与预算控制制度",
        "subjects": ("差旅报销", "招待费申请", "采购付款", "备用金核销", "发票复核"),
        "sections": ("预算申请", "票据要求", "金额权限", "超标处理", "版本说明"),
        "roles": ("财务复核人", "预算负责人"),
        "portals": ("财务共享平台", "费用控制中心", "电子报销系统"),
        "conditions": ("发票信息不完整", "费用超过预算", "付款对象发生变更"),
        "actions": ("冻结当前单据", "上传补充凭证"),
    },
    {
        "code": "SEC",
        "name": "信息安全",
        "title": "账号数据与安全事件制度",
        "subjects": ("特权账号开通", "数据外发", "终端遗失上报", "密码重置", "日志调阅"),
        "sections": ("身份认证", "数据保护", "事件响应", "审计留痕", "版本说明"),
        "roles": ("安全管理员", "数据所有者"),
        "portals": ("安全运营平台", "统一身份中心", "数据外发门户"),
        "conditions": ("发现异常登录", "设备疑似丢失", "敏感数据误发"),
        "actions": ("隔离相关账号", "提交安全事件单"),
    },
    {
        "code": "PROC",
        "name": "采购",
        "title": "供应商与采购执行规范",
        "subjects": ("供应商准入", "询价申请", "采购订单变更", "到货验收", "供应商复评"),
        "sections": ("供应商管理", "询比价流程", "订单执行", "异常采购", "版本说明"),
        "roles": ("采购经理", "需求部门负责人"),
        "portals": ("采购协同平台", "供应商门户", "订单管理中心"),
        "conditions": ("唯一供应商供货", "到货数量不一致", "交付日期发生延误"),
        "actions": ("暂停验收流程", "发起异常采购审批"),
    },
    {
        "code": "OPS",
        "name": "IT运维",
        "title": "信息系统运行维护手册",
        "subjects": ("生产发布", "数据库变更", "故障升级", "备份恢复", "监控告警"),
        "sections": ("变更管理", "发布控制", "故障响应", "备份策略", "版本说明"),
        "roles": ("值班负责人", "系统所有者"),
        "portals": ("运维工单平台", "变更控制台", "统一监控中心"),
        "conditions": ("生产服务不可用", "数据库校验失败", "监控连续触发告警"),
        "actions": ("启动应急预案", "保留故障现场日志"),
    },
    {
        "code": "PROD",
        "name": "产品",
        "title": "企业终端产品配置手册",
        "subjects": ("会议终端激活", "摄像头配网", "显示模块升级", "麦克风校准", "远程管理开通"),
        "sections": ("产品配置", "网络接入", "功能参数", "故障排查", "版本说明"),
        "roles": ("产品支持工程师", "客户成功经理"),
        "portals": ("设备管理云台", "产品服务中心", "远程运维门户"),
        "conditions": ("设备激活失败", "升级过程意外中断", "音频测试未通过"),
        "actions": ("恢复出厂配置", "导出设备诊断包"),
    },
    {
        "code": "CS",
        "name": "客户服务",
        "title": "客户服务与工单响应标准",
        "subjects": ("客户投诉", "服务升级", "退款审核", "现场支持", "满意度回访"),
        "sections": ("工单受理", "响应等级", "升级路径", "客户回访", "版本说明"),
        "roles": ("服务主管", "客户成功负责人"),
        "portals": ("客户服务台", "服务工单中心", "客户成功平台"),
        "conditions": ("客户重复投诉", "核心业务受到影响", "现场处理未解决"),
        "actions": ("提升工单优先级", "通知服务主管"),
    },
    {
        "code": "LEGAL",
        "name": "合同合规",
        "title": "合同审查与合规管理制度",
        "subjects": ("合同用印", "保密协议审查", "条款偏离审批", "合同归档", "合规咨询"),
        "sections": ("合同起草", "法律审查", "授权用印", "归档要求", "版本说明"),
        "roles": ("法务经理", "业务部门负责人"),
        "portals": ("合同生命周期平台", "法务协作中心", "电子用印系统"),
        "conditions": ("对方拒绝标准条款", "签约主体发生变化", "合同原件遗失"),
        "actions": ("暂停合同流转", "发起条款偏离审批"),
    },
    {
        "code": "LOG",
        "name": "物流仓储",
        "title": "仓储配送与库存管理规范",
        "subjects": ("入库验收", "库存盘点", "跨仓调拨", "异常出库", "退货入仓"),
        "sections": ("收货验收", "库存控制", "调拨管理", "异常处理", "版本说明"),
        "roles": ("仓储主管", "物流运营负责人"),
        "portals": ("仓储管理系统", "物流协同平台", "库存控制台"),
        "conditions": ("实物与系统数量不符", "包装出现明显破损", "运输单据缺失"),
        "actions": ("锁定相关库存", "创建仓储差异单"),
    },
    {
        "code": "RD",
        "name": "研发发布",
        "title": "研发质量与版本发布规范",
        "subjects": ("代码合并", "测试准入", "版本发布", "缺陷关闭", "开源组件引入"),
        "sections": ("研发流程", "质量门禁", "发布审批", "缺陷管理", "版本说明"),
        "roles": ("研发负责人", "质量负责人"),
        "portals": ("研发协作平台", "持续交付中心", "质量管理门户"),
        "conditions": ("自动化测试未通过", "线上出现阻断缺陷", "依赖组件存在高危漏洞"),
        "actions": ("停止发布流水线", "提交质量例外申请"),
    },
)


DEV_TARGETS = {
    "exact_term": 24,
    "semantic_paraphrase": 30,
    "numeric_table": 16,
    "parent_context": 14,
    "version_conflict": 10,
    "rewrite_needed": 6,
}
TEST_TARGETS = {
    "exact_term": 120,
    "semantic_paraphrase": 150,
    "numeric_table": 80,
    "parent_context": 70,
    "version_conflict": 50,
    "rewrite_needed": 30,
}


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def write_jsonl(path: Path, rows: list[dict]) -> None:
    path.write_text(
        "".join(json.dumps(row, ensure_ascii=False) + "\n" for row in rows),
        encoding="utf-8",
    )


def make_fact(domain: dict, global_doc: int, local_doc: int, fact_number: int, title: str) -> dict:
    category = CATEGORY_ORDER[fact_number % len(CATEGORY_ORDER)]
    clause_code = f"{domain['code']}-{global_doc + 1:02d}-{fact_number + 1:03d}"
    fact_id = f"FACT-{clause_code}"
    scope = SCOPES[local_doc]
    subject = f"{scope}{domain['subjects'][fact_number % len(domain['subjects'])]}"
    section = domain["sections"][(fact_number // 4) % len(domain["sections"])]
    role_primary, role_secondary = domain["roles"]
    portal = domain["portals"][(fact_number + local_doc) % len(domain["portals"])]
    condition = domain["conditions"][(fact_number + global_doc) % len(domain["conditions"])]
    action_primary, action_secondary = domain["actions"]
    version = f"V{2 + local_doc}.{1 + global_doc % 4}"
    effective = date(2026, 1, 1) + timedelta(days=17 * global_doc)

    if category == "exact_term":
        resource_code = f"{domain['code']}-{local_doc + 1:02d}-ZX{fact_number + 11:02d}"
        support = (
            f"条款 {clause_code}：专用业务标识 {resource_code} 仅用于{subject}。"
            f"申请人应通过{portal}发起流程，由{role_primary}进行首次确认；该标识不得用于其他事项。"
        )
        question = f"业务标识 {resource_code} 对应什么事项，应该从哪个入口办理？"
        answer = f"{resource_code}用于{subject}，办理入口是{portal}。"
        keywords = [resource_code, portal]
    elif category == "semantic_paraphrase":
        days = 1 + (global_doc * 3 + fact_number) % 12
        trigger = ("计划执行", "正式生效", "业务开始")[(global_doc + fact_number) % 3]
        support = (
            f"条款 {clause_code}：办理{subject}时，申请人必须在{trigger}前至少{days}个工作日"
            f"通过{portal}提交完整材料。提交后先由{role_primary}核验，再转交{role_secondary}备案。"
        )
        question = f"如果要办理{subject}，最晚需要提前几个工作日提交材料？"
        answer = f"至少提前{days}个工作日。"
        keywords = [str(days), "工作日"]
    elif category == "numeric_table":
        amount = 800 + ((global_doc * 137 + fact_number * 211) % 45) * 200
        support = (
            f"条款 {clause_code}：{subject}的单笔免追加审批上限为{amount}元。"
            f"金额不超过该上限时由{role_primary}复核；超过上限时必须增加{role_secondary}审批并上传预算依据。"
        )
        question = f"{subject}免于{role_secondary}追加审批的单笔金额上限是多少？"
        answer = f"单笔上限为{amount}元。"
        keywords = [str(amount), "元"]
    elif category == "parent_context":
        hours = 1 + (global_doc + fact_number * 2) % 12
        support = (
            f"条款 {clause_code}：当{condition}时，{subject}不得继续按普通流程处理。"
            f"经办人应先{action_primary}，并在{hours}小时内{action_secondary}；"
            f"完成上述两步后，才可由{role_secondary}决定是否恢复原流程。"
        )
        question = f"遇到{condition}时，{subject}应当按什么顺序处理？"
        answer = f"先{action_primary}，并在{hours}小时内{action_secondary}。"
        keywords = [action_primary, action_secondary, str(hours)]
    elif category == "version_conflict":
        hours = 4 + (global_doc * 2 + fact_number) % 44
        old_hours = hours + 12
        support = (
            f"条款 {clause_code}：本文件 {version} 自{effective.isoformat()}起生效。"
            f"针对{subject}，当前版本规定必须在{hours}小时内完成处理；旧版曾规定为{old_hours}小时，"
            f"两者冲突时一律以 {version} 为准。"
        )
        question = f"按照《{title}》{version}，{subject}必须在多少小时内处理？"
        answer = f"必须在{hours}小时内处理。"
        keywords = [str(hours), "小时"]
    else:
        colloquial = ("补手续", "走特批", "开权限", "提工单")[(global_doc + fact_number) % 4]
        formal = f"{subject}专项流程"
        support = (
            f"条款 {clause_code}：员工口头所称的“{colloquial}”，在本制度中的正式名称为“{formal}”。"
            f"该事项只能从{portal}进入，选择{clause_code}对应模板后提交，线下消息不能替代正式记录。"
        )
        question = (
            f"关于{subject}，{scope}同事口头说要“{colloquial}”，"
            "实际应从哪个系统入口走正式流程？"
        )
        answer = f"应通过{portal}进入{formal}。"
        keywords = [portal, formal]

    return {
        "fact_id": fact_id,
        "clause_code": clause_code,
        "domain_code": domain["code"],
        "domain": domain["name"],
        "document_key": f"DOC-{global_doc + 1:03d}",
        "document_title": title,
        "section": section,
        "category": category,
        "subject": subject,
        "support_text": support,
        "question": question,
        "reference_answer": answer,
        "answer_keywords": keywords,
        "version": version,
    }


def render_markdown(path: Path, title: str, facts: list[dict]) -> None:
    lines = [f"# {title}", "", "本文件为 Agentic RAG 检索评测生成的虚构企业资料。", ""]
    current = None
    for fact in facts:
        if fact["section"] != current:
            current = fact["section"]
            lines.extend([f"## {current}", ""])
        lines.extend([f"### {fact['clause_code']}", "", fact["support_text"], ""])
    path.write_text("\n".join(lines), encoding="utf-8")


def render_text(path: Path, title: str, facts: list[dict]) -> None:
    lines = [title, "本文件为 Agentic RAG 检索评测生成的虚构企业资料。", ""]
    current = None
    for fact in facts:
        if fact["section"] != current:
            current = fact["section"]
            lines.extend([f"【{current}】", ""])
        lines.extend([fact["support_text"], ""])
    path.write_text("\n".join(lines), encoding="utf-8")


def render_docx(path: Path, title: str, facts: list[dict]) -> None:
    document = Document()
    document.add_heading(title, 0)
    document.add_paragraph("本文件为 Agentic RAG 检索评测生成的虚构企业资料。")
    current = None
    for fact in facts:
        if fact["section"] != current:
            current = fact["section"]
            document.add_heading(current, level=1)
        document.add_heading(fact["clause_code"], level=2)
        document.add_paragraph(fact["support_text"])
    document.save(path)


def render_pdf(path: Path, title: str, facts: list[dict]) -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "BenchmarkTitle",
        parent=styles["Title"],
        fontName="STSong-Light",
        fontSize=18,
        alignment=TA_CENTER,
    )
    heading_style = ParagraphStyle(
        "BenchmarkHeading",
        parent=styles["Heading2"],
        fontName="STSong-Light",
        fontSize=12,
        leading=18,
    )
    body_style = ParagraphStyle(
        "BenchmarkBody",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=10,
        leading=16,
    )
    story = [
        Paragraph(title, title_style),
        Spacer(1, 5 * mm),
        Paragraph("本文件为 Agentic RAG 检索评测生成的虚构企业资料。", body_style),
    ]
    for index, fact in enumerate(facts):
        story.append(Paragraph(f"{fact['section']}｜{fact['clause_code']}", heading_style))
        story.append(Paragraph(fact["support_text"], body_style))
        story.append(Spacer(1, 3 * mm))
        if (index + 1) % 4 == 0 and index + 1 < len(facts):
            story.append(PageBreak())
    SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    ).build(story)


def render_xlsx(path: Path, title: str, facts: list[dict]) -> None:
    workbook = Workbook()
    workbook.remove(workbook.active)
    grouped: dict[str, list[dict]] = defaultdict(list)
    for fact in facts:
        grouped[fact["section"]].append(fact)
    for number, (section, section_facts) in enumerate(grouped.items(), start=1):
        sheet = workbook.create_sheet(f"{number}-{section}"[:31])
        sheet.append(["条款编号", "适用事项", "版本", "条款正文"])
        for fact in section_facts:
            sheet.append(
                [fact["clause_code"], fact["subject"], fact["version"], fact["support_text"]]
            )
    note = workbook.create_sheet("文档说明")
    note.append(["文档标题", title])
    note.append(["数据声明", "本文件全部内容均为虚构评测数据"])
    workbook.save(path)


def render_document(path: Path, title: str, facts: list[dict]) -> None:
    renderer = {
        ".md": render_markdown,
        ".txt": render_text,
        ".docx": render_docx,
        ".pdf": render_pdf,
        ".xlsx": render_xlsx,
    }[path.suffix]
    renderer(path, title, facts)


def qa_row(fact: dict, split: str, number: int) -> dict:
    return {
        "id": f"{split.upper()}-{number:04d}",
        "question": fact["question"],
        "reference_answer": fact["reference_answer"],
        "answer_keywords": fact["answer_keywords"],
        "answerable": True,
        "category": fact["category"],
        "domain_code": fact["domain_code"],
        "difficulty": "hard" if fact["category"] in {"version_conflict", "rewrite_needed"} else "medium",
        "relevant_fact_ids": [fact["fact_id"]],
        "relevant_clause_codes": [fact["clause_code"]],
        "relevant_source": fact["source"],
        "relevant_child_ids": [],
        "relevant_parent_ids": [],
    }


def select_split(facts: list[dict], targets: dict[str, int], split: str, rng: random.Random) -> list[dict]:
    by_category: dict[str, list[dict]] = defaultdict(list)
    for fact in facts:
        by_category[fact["category"]].append(fact)
    selected: list[dict] = []
    for category, count in targets.items():
        rng.shuffle(by_category[category])
        if len(by_category[category]) < count:
            raise RuntimeError(f"not enough facts for {category}")
        chosen = by_category[category][:count]
        selected.extend(chosen)
        chosen_ids = {fact["fact_id"] for fact in chosen}
        facts[:] = [fact for fact in facts if fact["fact_id"] not in chosen_ids]
    rng.shuffle(selected)
    return [qa_row(fact, split, number) for number, fact in enumerate(selected, start=1)]


def generate_grounding_set(test_rows: list[dict], rng: random.Random) -> list[dict]:
    eligible = [row for row in test_rows if row["domain_code"] != "PROD"]
    rng.shuffle(eligible)
    answerable = []
    for number, row in enumerate(eligible[:100], start=1):
        answerable.append({**row, "id": f"GROUND-A-{number:03d}"})

    unanswerable = []
    missing_items = ("量子密钥轮换", "宠物医疗补贴", "火星差旅氧气费", "水下设备保修", "无人机通勤津贴")
    for number in range(1, 101):
        domain = DOMAINS[(number - 1) % len(DOMAINS)]
        missing = missing_items[(number - 1) % len(missing_items)]
        unanswerable.append(
            {
                "id": f"GROUND-U-{number:03d}",
                "question": (
                    f"根据{domain['name']}制度，虚构编号 VOID-{number:03d} 对应的{missing}"
                    "必须在多少小时内办理？"
                ),
                "reference_answer": "知识库中没有相关规定",
                "answer_keywords": [],
                "answerable": False,
                "category": "unanswerable",
                "domain_code": domain["code"],
                "difficulty": "hard",
                "relevant_fact_ids": [],
                "relevant_clause_codes": [],
                "relevant_source": "",
                "relevant_child_ids": [],
                "relevant_parent_ids": [],
            }
        )
    rows = answerable + unanswerable
    rng.shuffle(rows)
    return rows


def main() -> None:
    rng = random.Random(SEED)
    BENCHMARK_DIR.mkdir(parents=True, exist_ok=True)
    CORPUS_DIR.mkdir(parents=True, exist_ok=True)
    for old in CORPUS_DIR.iterdir():
        if old.is_file() and old.suffix.lower() in FORMATS:
            old.unlink()

    all_facts: list[dict] = []
    documents: list[dict] = []
    global_doc = 0
    for domain in DOMAINS:
        for local_doc, scope in enumerate(SCOPES):
            title = f"{scope}{domain['title']}（模拟评测版）"
            suffix = FORMATS[global_doc % len(FORMATS)]
            source = f"{global_doc + 1:03d}_{domain['code']}_{scope}_{domain['title']}{suffix}"
            facts = [
                make_fact(domain, global_doc, local_doc, fact_number, title)
                for fact_number in range(20)
            ]
            for fact in facts:
                fact["source"] = source
            path = CORPUS_DIR / source
            render_document(path, title, facts)
            documents.append(
                {
                    "document_key": f"DOC-{global_doc + 1:03d}",
                    "domain_code": domain["code"],
                    "domain": domain["name"],
                    "title": title,
                    "source": source,
                    "format": suffix,
                    "facts": len(facts),
                    "sha256": file_sha256(path),
                    "bytes": path.stat().st_size,
                }
            )
            all_facts.extend(facts)
            global_doc += 1

    remaining = list(all_facts)
    dev_rows = select_split(remaining, DEV_TARGETS, "dev", rng)
    test_rows = select_split(remaining, TEST_TARGETS, "test", rng)
    grounding_rows = generate_grounding_set(test_rows, rng)

    write_jsonl(BENCHMARK_DIR / "facts.jsonl", all_facts)
    write_jsonl(BENCHMARK_DIR / "retrieval_dev.jsonl", dev_rows)
    write_jsonl(BENCHMARK_DIR / "retrieval_test.jsonl", test_rows)
    write_jsonl(BENCHMARK_DIR / "grounding_test.jsonl", grounding_rows)

    manifest = {
        "benchmark": "enterprise_agentic_rag_benchmark_v1",
        "seed": SEED,
        "generated_at": date.today().isoformat(),
        "synthetic_data": True,
        "documents": documents,
        "counts": {
            "documents": len(documents),
            "facts": len(all_facts),
            "retrieval_dev": len(dev_rows),
            "retrieval_test": len(test_rows),
            "grounding_test": len(grounding_rows),
            "grounding_answerable": sum(row["answerable"] for row in grounding_rows),
            "grounding_unanswerable": sum(not row["answerable"] for row in grounding_rows),
        },
        "retrieval_dev_categories": dict(Counter(row["category"] for row in dev_rows)),
        "retrieval_test_categories": dict(Counter(row["category"] for row in test_rows)),
        "label_status": "fact_ids_only_until_link_ground_truth",
    }
    (BENCHMARK_DIR / "corpus_manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(json.dumps(manifest["counts"], ensure_ascii=False, indent=2))
    print(f"Benchmark: {BENCHMARK_DIR}")


if __name__ == "__main__":
    main()
